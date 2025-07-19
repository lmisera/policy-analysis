import requests
import io
import os
import re
import time
from PyPDF2 import PdfReader
from openai import OpenAI
from dotenv import load_dotenv
from serpapi import GoogleSearch
from docx import Document
from docx.shared import Pt

# Load environment variables
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Step 1: Download PDF from S3
def download_pdf_from_s3(s3_url):
    print("📥 Downloading PDF from S3...")
    response = requests.get(s3_url)
    response.raise_for_status()
    return io.BytesIO(response.content)

# Step 2: Extract text from PDF
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n\n"
    return text

# Step 3: Chunk the text
def chunk_text(text, max_chunk_length=3000):
    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) < max_chunk_length:
            current_chunk += para + "\n\n"
        else:
            chunks.append(current_chunk.strip())
            current_chunk = para + "\n\n"
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

# Step 4: Extract legal references
def extract_references(text):
    patterns = [
        r"(Section\s\d{3,5}.*?Public Law \d{3}-\d{3})",
        r"(CFR\s?\d{1,3}\s?(Part|§)?\s?\d+)",
        r"Section\s\d+[A-Z]?"
    ]
    matches = []
    for pattern in patterns:
        matches.extend(re.findall(pattern, text))
    return [m if isinstance(m, str) else m[0] for m in matches]

# Step 5: SerpAPI legal reference search
def search_reference_with_serpapi(query):
    params = {
        "engine": "google",
        "q": query,
        "api_key": os.getenv("SERPAPI_API_KEY")
    }
    search = GoogleSearch(params)
    results = search.get_dict()
    snippets = []
    if "organic_results" in results:
        for result in results["organic_results"][:3]:
            snippet = result.get("snippet", "")
            link = result.get("link", "")
            snippets.append(f"{snippet} [source: {link}]")
    return "\n".join(snippets) if snippets else "No relevant information found."

# Step 6: Summarize with context
def summarize_chunk_with_context(client, chunk_text, external_context):
    prompt = f"""
You are a policy analyst focused on nuclear energy. 

Summarize the following legislative text with an emphasis on anything relevant to nuclear energy. 
If the external references below clarify or affect interpretation, include them.

Legislative Text:
\"\"\"{chunk_text}\"\"\"

External Legal References:
\"\"\"{external_context}\"\"\"

Be concise and focus on aspects that might affect nuclear development, funding, tax credits, safety, regulation, or deployment.
Without exception, include any changes to things like the Inflation Reduction Act or anything related to nuclear fuel supply chains.
This should be seven pages long total — so each chunk should be detailed but not overly repetitive.
"""
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return response.choices[0].message.content.strip()

# Step 7: Condense all summaries
def summarize_all_summaries(summaries, client):
    joined_summary = "\n\n".join(summaries)
    prompt = (
        "You are a policy analyst. Summarize the following nuclear-related legislative summaries into a cohesive, "
        "concise 7-page report. Explicitly mention anything related to nuclear fuel supply chain, tax credits, regulatory amendments, or changes to the Inflation Reduction Act. "
        "Group similar themes, eliminate redundancy, and write clearly:\n\n"
        f"{joined_summary}"
    )
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

# Step 8: Export to Word doc
def export_summaries_to_docx(summaries, output_filename="Nuclear_Policy_Summary.docx"):
    doc = Document()
    doc.add_heading("Nuclear-Relevant Summary of Legislation", 0)
    doc.add_paragraph("This summary aggregates nuclear-related information from the legislative text and referenced legal materials.").italic = True
    doc.add_paragraph(" ")
    for i, summary in enumerate(summaries):
        doc.add_heading(f"Section {i+1}", level=1)
        para = doc.add_paragraph(summary)
        para.style.font.size = Pt(11)
    doc.save(output_filename)
    print(f"📄 Report saved as: {output_filename}")

# === Main ===
if __name__ == "__main__":
    s3_url = "https://my-mapbox-data-geojson.s3.us-east-2.amazonaws.com/BBB_HR1.pdf"
    pdf_file = download_pdf_from_s3(s3_url)
    full_text = extract_text_from_pdf(pdf_file)

    MAX_CHUNKS = 10  # Adjust based on token constraints
    chunks = chunk_text(full_text)[:MAX_CHUNKS]
    print(f"✅ Split into {len(chunks)} chunks.")

    summaries = []
    for i, chunk in enumerate(chunks):
        print(f"✏️ Summarizing chunk {i+1}/{len(chunks)}...")
        refs = extract_references(chunk)
        extra_context = ""
        for ref in refs:
            print(f"🔎 Looking up: {ref}")
            search_result = search_reference_with_serpapi(ref)
            extra_context += f"\nContext for {ref}:\n{search_result}\n"
            time.sleep(1)  # Avoid hitting rate limits
        summary = summarize_chunk_with_context(client, chunk, extra_context)
        summaries.append(summary)

    final_summary = summarize_all_summaries(summaries, client)
    export_summaries_to_docx([final_summary])